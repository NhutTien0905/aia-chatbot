"""Generate sample insurance documents for testing: 1 DOCX + 1 PDF."""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import os

OUTPUT_DIR = "sample_docs"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# =============================================================================
# 1. DOCX - Hợp đồng bảo hiểm nhân thọ
# =============================================================================

def create_life_insurance_docx():
    doc = Document()

    # Title
    title = doc.add_heading("HỢP ĐỒNG BẢO HIỂM NHÂN THỌ", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Số hợp đồng: LI-2024-001234")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Section 1
    doc.add_heading("Điều 1: Các bên tham gia hợp đồng", level=1)
    doc.add_paragraph(
        "1.1. Bên bảo hiểm (Công ty bảo hiểm): Công ty TNHH Bảo hiểm An Phát Việt Nam, "
        "địa chỉ: Tầng 15, Tòa nhà Landmark 72, Phạm Hùng, Nam Từ Liêm, Hà Nội. "
        "Giấy phép kinh doanh số: 0101234567, cấp ngày 15/03/2010."
    )
    doc.add_paragraph(
        "1.2. Bên mua bảo hiểm (Chủ hợp đồng): Nguyễn Văn An, sinh ngày 15/06/1985, "
        "CCCD số: 001085012345, địa chỉ thường trú: 45 Nguyễn Trãi, Thanh Xuân, Hà Nội. "
        "Số điện thoại: 0912-345-678."
    )
    doc.add_paragraph(
        "1.3. Người được bảo hiểm: Nguyễn Văn An (đồng thời là chủ hợp đồng)."
    )
    doc.add_paragraph(
        "1.4. Người thụ hưởng: Trần Thị Bình (vợ), sinh ngày 22/09/1987, "
        "CCCD số: 001087098765."
    )

    # Section 2
    doc.add_heading("Điều 2: Phạm vi bảo hiểm", level=1)
    doc.add_paragraph(
        "2.1. Quyền lợi tử vong hoặc thương tật toàn bộ vĩnh viễn: "
        "Trong trường hợp Người được bảo hiểm tử vong hoặc bị thương tật toàn bộ vĩnh viễn "
        "trong thời hạn hợp đồng, Bên bảo hiểm sẽ chi trả 100% Số tiền bảo hiểm, "
        "tương đương 500.000.000 VNĐ (Năm trăm triệu đồng)."
    )
    doc.add_paragraph(
        "2.2. Quyền lợi bệnh hiểm nghèo: Nếu Người được bảo hiểm được chẩn đoán mắc một "
        "trong 34 bệnh hiểm nghèo theo danh mục tại Phụ lục A, Bên bảo hiểm sẽ chi trả "
        "trước 50% Số tiền bảo hiểm, tương đương 250.000.000 VNĐ."
    )
    doc.add_paragraph(
        "2.3. Quyền lợi đáo hạn: Khi hợp đồng đáo hạn và Người được bảo hiểm còn sống, "
        "Bên bảo hiểm sẽ chi trả toàn bộ Giá trị tài khoản hợp đồng tại thời điểm đáo hạn, "
        "bao gồm phí bảo hiểm đã đóng và lãi tích lũy (lãi suất cam kết tối thiểu 3%/năm)."
    )
    doc.add_paragraph(
        "2.4. Quyền lợi tai nạn bổ sung: Trường hợp tử vong do tai nạn, "
        "Bên bảo hiểm chi trả thêm 100% Số tiền bảo hiểm (tổng chi trả: 1.000.000.000 VNĐ)."
    )

    # Section 3
    doc.add_heading("Điều 3: Phí bảo hiểm", level=1)
    doc.add_paragraph(
        "3.1. Phí bảo hiểm định kỳ: 12.500.000 VNĐ/quý (Mười hai triệu năm trăm nghìn đồng), "
        "tương đương 50.000.000 VNĐ/năm."
    )
    doc.add_paragraph(
        "3.2. Phương thức đóng phí: Chuyển khoản ngân hàng hoặc trích nợ tự động vào ngày 15 "
        "đầu mỗi quý."
    )
    doc.add_paragraph(
        "3.3. Thời hạn đóng phí: 20 năm kể từ ngày hợp đồng có hiệu lực."
    )
    doc.add_paragraph(
        "3.4. Thời gian gia hạn đóng phí: 60 ngày kể từ ngày đến hạn đóng phí. "
        "Trong thời gian gia hạn, hợp đồng vẫn có hiệu lực."
    )

    # Section 4
    doc.add_heading("Điều 4: Các trường hợp loại trừ bảo hiểm", level=1)
    doc.add_paragraph(
        "Bên bảo hiểm không chịu trách nhiệm chi trả quyền lợi bảo hiểm trong các trường hợp sau:"
    )
    doc.add_paragraph(
        "4.1. Người được bảo hiểm tự tử trong vòng 2 năm đầu kể từ ngày hợp đồng có hiệu lực "
        "hoặc ngày khôi phục hiệu lực hợp đồng gần nhất."
    )
    doc.add_paragraph(
        "4.2. Người được bảo hiểm chết do hành vi cố ý của Chủ hợp đồng hoặc Người thụ hưởng."
    )
    doc.add_paragraph(
        "4.3. Người được bảo hiểm tham gia các hoạt động nguy hiểm: đua xe, nhảy dù, "
        "leo núi trên 4000m, lặn biển sâu trên 30m mà không có chứng chỉ chuyên nghiệp."
    )
    doc.add_paragraph(
        "4.4. Tử vong hoặc thương tật do chiến tranh, bạo loạn, khủng bố, "
        "hoặc nhiễm phóng xạ hạt nhân."
    )
    doc.add_paragraph(
        "4.5. Người được bảo hiểm bị ảnh hưởng bởi rượu (nồng độ cồn trong máu vượt 80mg/dl) "
        "hoặc sử dụng chất ma túy, chất kích thích trái phép tại thời điểm xảy ra sự kiện bảo hiểm."
    )

    # Section 5
    doc.add_heading("Điều 5: Thời hạn hợp đồng", level=1)
    doc.add_paragraph(
        "5.1. Ngày hiệu lực: 01/01/2024."
    )
    doc.add_paragraph(
        "5.2. Ngày đáo hạn: 01/01/2054 (thời hạn 30 năm)."
    )
    doc.add_paragraph(
        "5.3. Tuổi tham gia bảo hiểm: 38 tuổi (tính theo năm sinh)."
    )

    # Section 6
    doc.add_heading("Điều 6: Quyền và nghĩa vụ của các bên", level=1)
    doc.add_paragraph(
        "6.1. Quyền của Chủ hợp đồng: Thay đổi Người thụ hưởng, yêu cầu giảm Số tiền bảo hiểm, "
        "tạm ứng từ Giá trị hoàn lại (tối đa 70%), chuyển nhượng hợp đồng."
    )
    doc.add_paragraph(
        "6.2. Nghĩa vụ của Chủ hợp đồng: Đóng phí đúng hạn, khai báo trung thực tình trạng "
        "sức khỏe, thông báo cho Bên bảo hiểm khi có thay đổi nghề nghiệp hoặc địa chỉ."
    )
    doc.add_paragraph(
        "6.3. Quyền của Bên bảo hiểm: Từ chối chi trả nếu phát hiện gian lận, "
        "điều chỉnh phí bảo hiểm theo quy định pháp luật."
    )
    doc.add_paragraph(
        "6.4. Nghĩa vụ của Bên bảo hiểm: Chi trả quyền lợi trong vòng 15 ngày làm việc "
        "kể từ khi nhận đủ hồ sơ yêu cầu hợp lệ."
    )

    # Section 7
    doc.add_heading("Điều 7: Thủ tục yêu cầu giải quyết quyền lợi", level=1)
    doc.add_paragraph(
        "7.1. Hồ sơ yêu cầu bao gồm: Đơn yêu cầu giải quyết quyền lợi (theo mẫu), "
        "bản sao CCCD của người yêu cầu, giấy chứng tử (trường hợp tử vong), "
        "hồ sơ bệnh án (trường hợp bệnh hiểm nghèo), biên bản tai nạn (nếu có)."
    )
    doc.add_paragraph(
        "7.2. Thời hạn nộp hồ sơ: Trong vòng 180 ngày kể từ ngày xảy ra sự kiện bảo hiểm."
    )
    doc.add_paragraph(
        "7.3. Phương thức chi trả: Chuyển khoản vào tài khoản ngân hàng của Người thụ hưởng "
        "đã đăng ký."
    )

    filepath = os.path.join(OUTPUT_DIR, "Hop_Dong_Bao_Hiem_Nhan_Tho_LI2024001234.docx")
    doc.save(filepath)
    print(f"✓ Created: {filepath}")


# =============================================================================
# 2. PDF - Hợp đồng bảo hiểm xe ô tô
# =============================================================================

def create_auto_insurance_pdf():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # Add Unicode font
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    bold_font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"

    if not os.path.exists(font_path):
        # Fallback: try another common location
        font_path = "/usr/share/fonts/TTF/DejaVuSans.ttf"
        bold_font_path = "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf"

    pdf.add_font("DejaVu", "", font_path, uni=True)
    pdf.add_font("DejaVu", "B", bold_font_path, uni=True)

    # --- Page 1 ---
    pdf.add_page()
    pdf.set_font("DejaVu", "B", 16)
    pdf.cell(0, 12, "HỢP ĐỒNG BẢO HIỂM XE Ô TÔ", ln=True, align="C")
    pdf.set_font("DejaVu", "B", 11)
    pdf.cell(0, 8, "Số hợp đồng: AUTO-2024-005678", ln=True, align="C")
    pdf.cell(0, 8, "Ngày cấp: 15/03/2024", ln=True, align="C")
    pdf.ln(8)

    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 1: Thông tin các bên", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "1.1. Bên bảo hiểm: Tổng Công ty Bảo hiểm Toàn Cầu (GlobalCare Insurance), "
        "Trụ sở: 88 Lê Lợi, Quận 1, TP. Hồ Chí Minh. MST: 0301234567.\n\n"
        "1.2. Bên mua bảo hiểm: Trần Minh Đức, sinh ngày 20/11/1990, "
        "CCCD: 079090123456, địa chỉ: 123 Nguyễn Huệ, Quận 1, TP. Hồ Chí Minh. "
        "SĐT: 0908-123-456.\n\n"
        "1.3. Thông tin xe được bảo hiểm:\n"
        "   - Nhãn hiệu: Toyota Camry 2.5Q\n"
        "   - Năm sản xuất: 2023\n"
        "   - Biển số: 51A-123.45\n"
        "   - Số khung: JTDKN3DU5A0123456\n"
        "   - Số máy: 2AR-FE-7891011\n"
        "   - Màu sơn: Đen\n"
        "   - Giá trị xe: 1.200.000.000 VNĐ"
    )
    pdf.ln(4)

    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 2: Phạm vi bảo hiểm vật chất xe", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "2.1. Bảo hiểm thiệt hại vật chất xe (Comprehensive Coverage):\n"
        "Bên bảo hiểm bồi thường cho thiệt hại vật chất xảy ra đối với xe được bảo hiểm do:\n"
        "   a) Tai nạn, va chạm, lật xe, rơi xuống vực\n"
        "   b) Cháy, nổ (kể cả cháy do chập điện)\n"
        "   c) Thiên tai: bão, lũ, sét đánh, mưa đá, động đất, sóng thần\n"
        "   d) Hành vi phá hoại của bên thứ ba\n"
        "   e) Mất cắp toàn bộ xe\n\n"
        "2.2. Mức bồi thường tối đa: 100% giá trị xe tại thời điểm xảy ra tổn thất, "
        "nhưng không vượt quá 1.200.000.000 VNĐ.\n\n"
        "2.3. Mức khấu trừ (Deductible): 500.000 VNĐ/vụ đối với thiệt hại vật chất, "
        "20% giá trị bồi thường đối với mất cắp bộ phận."
    )
    pdf.ln(4)

    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 3: Bảo hiểm trách nhiệm dân sự", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "3.1. Trách nhiệm dân sự đối với người thứ ba:\n"
        "   - Thiệt hại về người: tối đa 150.000.000 VNĐ/người/vụ\n"
        "   - Thiệt hại về tài sản: tối đa 100.000.000 VNĐ/vụ\n\n"
        "3.2. Trách nhiệm dân sự đối với hành khách trên xe:\n"
        "   - Số chỗ ngồi được bảo hiểm: 5 chỗ\n"
        "   - Mức trách nhiệm: 50.000.000 VNĐ/người/vụ\n\n"
        "3.3. Bảo hiểm tai nạn lái xe và phụ xe:\n"
        "   - Tử vong/thương tật vĩnh viễn: 100.000.000 VNĐ/người\n"
        "   - Chi phí y tế: 20.000.000 VNĐ/người/vụ"
    )

    # --- Page 2 ---
    pdf.add_page()
    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 4: Các trường hợp loại trừ", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "Bên bảo hiểm không bồi thường trong các trường hợp:\n\n"
        "4.1. Lái xe không có Giấy phép lái xe hợp lệ hoặc Giấy phép lái xe không phù hợp "
        "với loại xe được bảo hiểm.\n\n"
        "4.2. Lái xe có nồng độ cồn trong máu vượt quá 50mg/100ml hoặc trong hơi thở "
        "vượt quá 0.25mg/lít, hoặc sử dụng chất ma túy, chất kích thích bị cấm.\n\n"
        "4.3. Xe sử dụng để đua xe, thi đấu, chạy thử, hoặc kéo đẩy xe khác "
        "(trừ trường hợp xe được thiết kế để kéo).\n\n"
        "4.4. Thiệt hại do hao mòn tự nhiên, giảm giá trị theo thời gian, hỏng hóc cơ khí "
        "hoặc điện không phải do tai nạn.\n\n"
        "4.5. Thiệt hại xảy ra ngoài lãnh thổ Việt Nam.\n\n"
        "4.6. Xe chở hàng quá tải trọng cho phép trên 20% hoặc chở quá số người quy định.\n\n"
        "4.7. Thiệt hại do chiến tranh, đình công, bạo loạn, khủng bố, phóng xạ hạt nhân."
    )
    pdf.ln(4)

    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 5: Phí bảo hiểm", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "5.1. Phí bảo hiểm vật chất xe: 1.2% x 1.200.000.000 = 14.400.000 VNĐ/năm\n\n"
        "5.2. Phí bảo hiểm trách nhiệm dân sự bắt buộc: 530.000 VNĐ/năm\n\n"
        "5.3. Phí bảo hiểm trách nhiệm dân sự tự nguyện: 1.200.000 VNĐ/năm\n\n"
        "5.4. Phí bảo hiểm tai nạn lái phụ xe: 300.000 VNĐ/năm\n\n"
        "5.5. Tổng phí bảo hiểm: 16.430.000 VNĐ/năm (Mười sáu triệu bốn trăm ba mươi nghìn đồng)\n\n"
        "5.6. Phương thức thanh toán: Thanh toán một lần trước ngày hiệu lực hợp đồng.\n\n"
        "5.7. Giảm phí: Giảm 10% phí vật chất xe nếu không có khiếu nại trong năm trước "
        "(No-Claim Discount)."
    )
    pdf.ln(4)

    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 6: Thời hạn bảo hiểm", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "6.1. Ngày bắt đầu: 15/03/2024, 00:00 giờ\n\n"
        "6.2. Ngày kết thúc: 14/03/2025, 24:00 giờ\n\n"
        "6.3. Hợp đồng tự động gia hạn thêm 1 năm nếu không có bên nào thông báo "
        "chấm dứt trước 30 ngày."
    )

    # --- Page 3 ---
    pdf.add_page()
    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 7: Thủ tục yêu cầu bồi thường", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "7.1. Khi xảy ra tai nạn hoặc tổn thất, Bên mua bảo hiểm phải:\n"
        "   a) Thông báo cho Bên bảo hiểm trong vòng 24 giờ qua hotline: 1900-1234\n"
        "   b) Giữ nguyên hiện trường (nếu có thể) và chụp ảnh thiệt hại\n"
        "   c) Trình báo cơ quan công an (đối với tai nạn có thiệt hại người hoặc tài sản lớn)\n\n"
        "7.2. Hồ sơ yêu cầu bồi thường bao gồm:\n"
        "   - Đơn yêu cầu bồi thường (theo mẫu)\n"
        "   - Bản sao Giấy phép lái xe, Đăng ký xe\n"
        "   - Biên bản tai nạn giao thông (nếu có)\n"
        "   - Hình ảnh thiệt hại\n"
        "   - Hóa đơn sửa chữa (bản gốc)\n"
        "   - Bản sao CCCD người yêu cầu\n\n"
        "7.3. Thời hạn giải quyết: 15 ngày làm việc kể từ khi nhận đủ hồ sơ hợp lệ. "
        "Trường hợp phức tạp cần giám định: tối đa 30 ngày làm việc.\n\n"
        "7.4. Phương thức bồi thường:\n"
        "   - Sửa chữa tại garage ủy quyền (không cần ứng trước)\n"
        "   - Hoặc bồi thường bằng tiền mặt theo hóa đơn sửa chữa thực tế"
    )
    pdf.ln(4)

    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 8: Quyền và nghĩa vụ các bên", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "8.1. Nghĩa vụ của Bên mua bảo hiểm:\n"
        "   - Cung cấp thông tin trung thực về xe và lịch sử tai nạn\n"
        "   - Bảo dưỡng xe định kỳ theo khuyến cáo nhà sản xuất\n"
        "   - Thông báo khi có thay đổi mục đích sử dụng xe\n"
        "   - Áp dụng các biện pháp phòng ngừa tổn thất hợp lý\n\n"
        "8.2. Quyền của Bên mua bảo hiểm:\n"
        "   - Yêu cầu bồi thường khi xảy ra sự kiện bảo hiểm\n"
        "   - Chấm dứt hợp đồng trước hạn và nhận hoàn phí theo tỷ lệ\n"
        "   - Khiếu nại quyết định từ chối bồi thường\n\n"
        "8.3. Quyền của Bên bảo hiểm:\n"
        "   - Giám định thiệt hại trước khi bồi thường\n"
        "   - Từ chối bồi thường nếu phát hiện gian lận\n"
        "   - Truy đòi bên thứ ba gây ra thiệt hại (quyền thế quyền)\n\n"
        "8.4. Nghĩa vụ của Bên bảo hiểm:\n"
        "   - Cấp Giấy chứng nhận bảo hiểm\n"
        "   - Giải quyết bồi thường đúng thời hạn\n"
        "   - Bảo mật thông tin khách hàng"
    )
    pdf.ln(4)

    pdf.set_font("DejaVu", "B", 13)
    pdf.cell(0, 10, "Điều 9: Điều khoản chung", ln=True)
    pdf.set_font("DejaVu", "", 10)
    pdf.multi_cell(0, 6,
        "9.1. Mọi tranh chấp phát sinh từ hợp đồng này được giải quyết bằng thương lượng. "
        "Nếu không thương lượng được, các bên đưa ra Trung tâm Trọng tài Quốc tế Việt Nam "
        "(VIAC) để giải quyết.\n\n"
        "9.2. Hợp đồng này được lập thành 02 bản có giá trị pháp lý như nhau, "
        "mỗi bên giữ 01 bản.\n\n"
        "9.3. Hợp đồng có hiệu lực kể từ ngày Bên mua bảo hiểm đóng đủ phí bảo hiểm.\n\n"
        "9.4. Luật áp dụng: Pháp luật nước Cộng hòa Xã hội Chủ nghĩa Việt Nam."
    )

    filepath = os.path.join(OUTPUT_DIR, "Hop_Dong_Bao_Hiem_Xe_Oto_AUTO2024005678.pdf")
    pdf.output(filepath)
    print(f"✓ Created: {filepath}")


if __name__ == "__main__":
    print("Generating sample insurance documents...\n")
    create_life_insurance_docx()
    create_auto_insurance_pdf()
    print(f"\n✅ Done! Files saved in '{OUTPUT_DIR}/' folder.")
